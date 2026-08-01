from pathlib import Path
from flask import current_app
import markdown
from utils.logger import logger


class ReportGenerator:
    """Generates PDF, DOCX, Markdown, and HTML reports for ATS Analysis."""

    @classmethod
    def generate_pdf(cls, report_obj, resume_obj, user_obj, output_path: Path) -> Path:
        """Generate PDF report using ReportLab."""
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.lib import colors
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

            doc = SimpleDocTemplate(
                str(output_path),
                pagesize=letter,
                rightMargin=36,
                leftMargin=36,
                topMargin=36,
                bottomMargin=36
            )
            story = []
            styles = getSampleStyleSheet()

            # Custom Styles
            title_style = ParagraphStyle(
                'DocTitle',
                parent=styles['Heading1'],
                fontSize=22,
                leading=26,
                textColor=colors.HexColor('#0F172A'),
                spaceAfter=10
            )

            heading2_style = ParagraphStyle(
                'SectionHeading',
                parent=styles['Heading2'],
                fontSize=14,
                leading=18,
                textColor=colors.HexColor('#2563EB'),
                spaceBefore=12,
                spaceAfter=6
            )

            normal_style = ParagraphStyle(
                'Body',
                parent=styles['Normal'],
                fontSize=10,
                leading=14,
                textColor=colors.HexColor('#334155')
            )

            import html
            # Header
            story.append(Paragraph("AI Resume Analysis & ATS Optimization Report", title_style))
            c_name = html.escape(str(user_obj.name or 'Candidate'))
            c_email = html.escape(str(user_obj.email or ''))
            c_filename = html.escape(str(resume_obj.filename or ''))
            story.append(Paragraph(f"Candidate: <b>{c_name}</b> ({c_email}) | File: {c_filename}", normal_style))
            story.append(Spacer(1, 10))
            story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor('#CBD5E1'), spaceAfter=15))

            # Scores Summary Table
            score_data = [
                ["Overall ATS Score", f"{report_obj.ats_score}%"],
                ["Formatting Score", f"{report_obj.formatting_score}%"],
                ["Skills Score", f"{report_obj.skills_score}%"],
                ["Experience Score", f"{report_obj.experience_score}%"],
                ["Keywords & Action Verbs", f"{report_obj.keywords_score}%"],
                ["Readability", f"{report_obj.readability_score}%"]
            ]
            
            table = Table(score_data, colWidths=[240, 100])
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1E293B')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                ('TOPPADDING', (0, 0), (-1, -1), 6),
                ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#F8FAFC')),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#E2E8F0')),
            ]))
            story.append(table)
            story.append(Spacer(1, 15))

            # Executive Summary
            story.append(Paragraph("Executive Summary", heading2_style))
            summary_text = report_obj.summary or report_obj.analysis_json.get("executive_summary", "Detailed candidate profile evaluation.")
            story.append(Paragraph(html.escape(str(summary_text)), normal_style))
            story.append(Spacer(1, 10))

            # Strengths & Weaknesses
            analysis = report_obj.analysis_json
            if analysis:
                story.append(Paragraph("Key Strengths", heading2_style))
                for s in analysis.get("key_strengths", []):
                    story.append(Paragraph(f"• {html.escape(str(s))}", normal_style))

                story.append(Spacer(1, 10))
                story.append(Paragraph("Recommended Improvements", heading2_style))
                for area in analysis.get("improvement_areas", []):
                    story.append(Paragraph(f"• {html.escape(str(area))}", normal_style))

            doc.build(story)
            return output_path
        except Exception as e:
            logger.error(f"Failed to generate PDF report: {e}")
            raise e

    @classmethod
    def generate_docx(cls, report_obj, resume_obj, user_obj, output_path: Path) -> Path:
        """Generate Word DOCX report using python-docx."""
        try:
            import docx
            doc = docx.Document()
            doc.add_heading("AI Resume Analysis & ATS Optimization Report", 0)
            doc.add_paragraph(f"Candidate: {user_obj.name} ({user_obj.email})")
            doc.add_paragraph(f"Resume File: {resume_obj.filename}")

            doc.add_heading("ATS Score Breakdown", level=1)
            t = doc.add_table(rows=1, cols=2)
            hdr_cells = t.rows[0].cells
            hdr_cells[0].text = 'Metric'
            hdr_cells[1].text = 'Score'

            scores = [
                ("Overall ATS Score", f"{report_obj.ats_score}%"),
                ("Formatting Score", f"{report_obj.formatting_score}%"),
                ("Skills Score", f"{report_obj.skills_score}%"),
                ("Experience Score", f"{report_obj.experience_score}%"),
                ("Keywords Score", f"{report_obj.keywords_score}%"),
                ("Readability Score", f"{report_obj.readability_score}%")
            ]
            for m, s in scores:
                row_cells = t.add_row().cells
                row_cells[0].text = m
                row_cells[1].text = s

            analysis = report_obj.analysis_json
            doc.add_heading("Executive Summary", level=1)
            doc.add_paragraph(report_obj.summary or analysis.get("executive_summary", ""))

            doc.add_heading("Key Strengths", level=1)
            for s in analysis.get("key_strengths", []):
                doc.add_paragraph(s, style='List Bullet')

            doc.add_heading("Recommended Improvements", level=1)
            for imp in analysis.get("improvement_areas", []):
                doc.add_paragraph(imp, style='List Bullet')

            doc.save(str(output_path))
            return output_path
        except Exception as e:
            logger.error(f"Failed to generate DOCX report: {e}")
            raise e

    @classmethod
    def generate_markdown(cls, report_obj, resume_obj, user_obj) -> str:
        """Generate Markdown string for report."""
        analysis = report_obj.analysis_json
        md = f"""# AI Resume Analysis & ATS Optimization Report

**Candidate:** {user_obj.name} ({user_obj.email})  
**Resume File:** {resume_obj.filename}  
**Date:** {report_obj.created_at.strftime('%Y-%m-%d %H:%M UTC')}

---

## 📊 ATS Score Breakdown

| Metric | Score |
| :--- | :--- |
| **Overall ATS Score** | **{report_obj.ats_score}%** |
| Formatting Score | {report_obj.formatting_score}% |
| Skills Score | {report_obj.skills_score}% |
| Experience Score | {report_obj.experience_score}% |
| Keywords Score | {report_obj.keywords_score}% |
| Readability Score | {report_obj.readability_score}% |

---

## 📝 Executive Summary
{report_obj.summary or analysis.get("executive_summary", "Solid technical resume.")}

---

## 💪 Key Strengths
"""
        for str_item in analysis.get("key_strengths", []):
            md += f"- {str_item}\n"

        md += "\n## 🚀 Recommended Improvements\n"
        for imp in analysis.get("improvement_areas", []):
            md += f"- {imp}\n"

        return md

    @classmethod
    def generate_html(cls, report_obj, resume_obj, user_obj) -> str:
        """Generate HTML representation of report."""
        md_content = cls.generate_markdown(report_obj, resume_obj, user_obj)
        return markdown.markdown(md_content, extensions=['tables'])
