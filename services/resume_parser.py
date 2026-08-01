import re
import json
import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path
from utils.logger import logger


class ResumeParser:
    """Enterprise ATS Resume Parser supporting PDF, DOCX, DOC, TXT, and Image Resumes (OCR)."""

    SECTION_MAP = {
        "summary": ["professional summary", "summary", "profile", "executive summary", "about me", "career summary", "personal profile", "overview"],
        "objective": ["career objective", "objective", "professional objective", "target role"],
        "education": ["education", "academic details", "academic background", "academic qualifications", "qualification", "qualifications", "academic profile", "educational qualification", "education & qualifications", "educational background"],
        "experience": ["work experience", "experience", "employment history", "professional experience", "work history", "professional background", "career history", "relevant experience"],
        "internship": ["internship", "internships", "internship experience", "industrial training", "summer internship"],
        "projects": ["projects", "academic projects", "personal projects", "key projects", "selected projects", "technical projects"],
        "skills": ["technical skills", "skills", "professional skills", "core competencies", "technologies", "tools & technologies", "skills & tools", "technical proficiencies", "technical expertise", "skills & expertise"],
        "programming_languages": ["programming languages", "coding languages", "languages & syntax"],
        "frameworks": ["frameworks", "web frameworks", "frameworks & libraries"],
        "libraries": ["libraries", "packages", "modules"],
        "databases": ["databases", "database management", "dbms", "data stores"],
        "developer_tools": ["developer tools", "tools", "development tools", "software & tools", "ide & tools"],
        "concepts": ["concepts", "technical concepts", "domain expertise", "methodologies", "core concepts"],
        "cloud": ["cloud", "cloud platforms", "cloud computing", "cloud services"],
        "devops": ["devops", "ci/cd & devops", "infrastructure"],
        "operating_systems": ["operating systems", "os", "platforms"],
        "soft_skills": ["soft skills", "interpersonal skills", "key skills", "managerial skills"],
        "certifications": ["certifications", "certificates", "licenses", "accreditations", "professional certifications"],
        "courses": ["courses", "online courses", "training", "trainings", "workshops"],
        "achievements": ["achievements", "awards", "honors", "accomplishments", "awards & achievements", "honors & awards"],
        "languages": ["languages", "language proficiency", "languages known", "spoken languages"],
        "volunteer": ["volunteer experience", "volunteering", "community service", "social work"],
        "leadership": ["leadership", "positions of responsibility", "extracurricular", "extracurricular activities"],
        "research": ["research", "research experience", "research work"],
        "publications": ["publications", "papers", "conference papers", "patents"],
        "interests": ["interests", "hobbies", "personal interests", "activities"],
        "references": ["references", "referees"]
    }

    SKILL_TAXONOMY = {
        "programming_languages": [
            "python", "java", "c++", "c#", "c", "javascript", "typescript", "html", "css", "html5", "css3",
            "sql", "go", "golang", "rust", "php", "ruby", "swift", "kotlin", "r", "scala", "dart", "shell", "bash", "powershell"
        ],
        "frameworks": [
            "react", "react.js", "reactjs", "next.js", "nextjs", "vue", "vue.js", "vuejs", "angular", "angularjs",
            "svelte", "redux", "express", "express.js", "flask", "django", "fastapi", "spring", "spring boot",
            "asp.net", ".net", ".net core", "laravel", "ruby on rails", "rails", "tailwind", "tailwindcss", "bootstrap"
        ],
        "libraries": [
            "pandas", "numpy", "scikit-learn", "tensorflow", "keras", "pytorch", "opencv", "nlp",
            "natural language processing", "llm", "large language models", "openai", "groq", "langchain",
            "hugging face", "jquery", "webpack", "vite", "sass", "less"
        ],
        "databases": [
            "postgresql", "postgres", "mysql", "sqlite", "mongodb", "redis", "cassandra", "elasticsearch",
            "dynamodb", "oracle", "sql server", "mssql", "firebase", "supabase", "mariadb", "neo4j"
        ],
        "developer_tools": [
            "git", "github", "gitlab", "bitbucket", "docker", "postman", "swagger", "vs code", "visual studio code",
            "jira", "confluence", "agile", "scrum", "kanban", "unit testing", "pytest", "jest", "cypress", "selenium"
        ],
        "concepts": [
            "responsive design", "responsive layout", "frontend development", "backend development", "full stack development",
            "rest api", "restful apis", "ui/ux", "user interface", "web security", "authentication", "authorization",
            "database design", "object oriented programming", "oop", "microservices", "data structures", "algorithms"
        ],
        "cloud": [
            "aws", "amazon web services", "azure", "gcp", "google cloud", "cloudformation", "terraform"
        ],
        "devops": [
            "linux", "unix", "ci/cd", "kubernetes", "k8s", "docker", "ansible", "jenkins", "github actions",
            "gitlab ci", "nginx", "apache", "helm"
        ],
        "operating_systems": [
            "windows", "linux", "macos", "unix", "android", "ios", "ubuntu", "debian", "centos"
        ],
        "soft_skills": [
            "communication", "leadership", "problem solving", "teamwork", "time management", "critical thinking",
            "adaptability", "analytical skills", "collaboration", "creativity"
        ]
    }

    # =========================================================================
    # MODULAR EXTRACTION ENGINE
    # =========================================================================

    @classmethod
    def extract_text(cls, filepath: str, file_type: str) -> str:
        text = ""
        path = Path(filepath)

        if not path.exists():
            raise FileNotFoundError(f"File not found: {filepath}")

        ext = file_type.lower().strip(".")

        if ext == "pdf":
            text = cls._extract_pdf(path)
        elif ext == "docx":
            text = cls._extract_docx(path)
        elif ext == "doc":
            raise ValueError("Legacy binary .doc files are not supported. Please save your file as .docx or .pdf.")
        elif ext == "txt":
            try:
                text = path.read_text(encoding="utf-8", errors="ignore")
            except Exception:
                text = path.read_text(encoding="latin1", errors="ignore")
        elif ext in ["png", "jpg", "jpeg"]:
            text = cls._extract_image_ocr(path)
        else:
            raise ValueError(f"Unsupported file format: {file_type}")

        cleaned_text = cls.clean_text(text)
        return cleaned_text

    @classmethod
    def clean_text(cls, text: str) -> str:
        if not text:
            return ""

        text = cls._fix_spaced_text(text)
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', text)

        cleaned_lines = []
        prev_empty = False

        for line in text.splitlines():
            clean = line.strip()

            if not clean or clean in ["|", "•", "-", "_", "+", "*", ":"]:
                continue

            if re.match(r'^[\|\s\-\_\+\=\:\.\,•]+$', clean) and not re.search(r'[a-zA-Z0-9]', clean):
                continue

            if clean.startswith("|") and clean.endswith("|") and len(clean) > 2:
                clean = clean[1:-1].strip()

            clean = re.sub(r'\s*\|\s*', ' • ', clean)
            clean = re.sub(r'(?:[ \t]*•[ \t]*)+', ' • ', clean).strip(' •')

            if not clean:
                if not prev_empty:
                    cleaned_lines.append("")
                    prev_empty = True
            else:
                cleaned_lines.append(clean)
                prev_empty = False

        return "\n".join(cleaned_lines).strip()

    @classmethod
    def normalize_heading(cls, line: str) -> str:
        clean = line.strip()
        clean = re.sub(r'^[•\-\*\>\✓\○\d\.\s]+', '', clean)
        clean = re.sub(r'[:\-\_\=\#\*]+$', '', clean).strip().lower()

        if not clean or len(clean) > 45:
            return ""

        for std_key, aliases in cls.SECTION_MAP.items():
            for alias in aliases:
                if clean == alias or clean == alias + "s" or clean == alias.rstrip("s"):
                    return std_key

        return ""

    @classmethod
    def extract_personal_information(cls, lines: list, text: str) -> dict:
        email = cls._extract_email(text)
        phone = cls._extract_phone(text)
        links = cls._extract_links(text)
        name = cls._extract_name(lines, email)
        location = cls._extract_location(text)
        title = cls._extract_professional_title(lines, text)

        dob = ""
        dob_match = re.search(r'(?:dob|date of birth|born)[:\s]*([0-9]{1,2}[\/\-\.][0-9]{1,2}[\/\-\.][0-9]{2,4}|[A-Za-z]+\s+[0-9]{1,2},\s*[0-9]{4})', text, re.IGNORECASE)
        if dob_match:
            dob = dob_match.group(1)

        nationality = ""
        nat_match = re.search(r'(?:nationality|citizenship)[:\s]*([A-Za-z]+)', text, re.IGNORECASE)
        if nat_match:
            nationality = nat_match.group(1).title()

        return {
            "name": name,
            "full_name": name,
            "professional_title": title,
            "current_position": title,
            "email": email,
            "phone": phone,
            "address": location,
            "location": location,
            "linkedin": links.get("linkedin", ""),
            "github": links.get("github", ""),
            "portfolio": links.get("portfolio", ""),
            "website": links.get("portfolio", ""),
            "dob": dob,
            "nationality": nationality
        }

    @classmethod
    def _sanitize_lines(cls, lines: list, personal_info: dict) -> list:
        clean = []
        name = personal_info.get("name", "").lower()
        email = personal_info.get("email", "").lower()
        phone_digits = re.sub(r'\D', '', personal_info.get("phone", ""))
        title = personal_info.get("professional_title", "").lower()

        job_title_kw = ["developer", "engineer", "intern", "self employed", "freelancer", "manager", "designer", "architect", "analyst"]

        for line in lines:
            l_str = line.strip()
            if not l_str:
                continue
            if email and email in l_str.lower():
                continue
            if phone_digits and len(phone_digits) >= 7 and phone_digits in re.sub(r'\D', '', l_str):
                continue
            if name and name in l_str.lower() and len(l_str.split()) <= 4:
                continue
            if 'linkedin.com' in l_str.lower() or 'github.com' in l_str.lower():
                continue

            # Strip job title from Education lines if present without a degree keyword
            is_edu_context = any(deg in l_str.lower() for deg in ["bachelor", "master", "b.sc", "b.tech", "m.sc", "degree", "diploma", "school", "college", "university", "institute", "ssce", "hsc", "cbse"])
            if any(kw in l_str.lower() for kw in job_title_kw) and not is_edu_context and ("education" in lines or "academic" in lines):
                continue

            clean.append(l_str)

        return clean

    @classmethod
    def extract_education(cls, lines: list) -> list:
        """Parses structured education items into unified objects, strictly excluding job titles."""
        if not lines:
            return []
        entries = []
        job_title_kw = ["developer", "engineer", "intern", "self employed", "freelancer", "manager"]

        # Filter lines to make sure no standalone job title line enters education
        clean_edu_lines = []
        for l in lines:
            if any(kw in l.lower() for kw in job_title_kw) and not any(deg in l.lower() for deg in ["bachelor", "master", "b.sc", "b.tech", "m.sc", "degree", "diploma", "ssce", "hsc"]):
                continue
            clean_edu_lines.append(l)

        full_text = "\n".join(clean_edu_lines)
        blocks = re.split(r'\n(?=[A-Z][a-zA-Z0-9\s,\.\(\)]*(?:Institute|University|College|School|Academy|Bachelor|Master|B\.Sc|M\.Sc|B\.Tech|M\.Tech|Diploma|Degree))', full_text)

        for block in blocks:
            b_lines = [l.strip() for l in block.splitlines() if l.strip()]
            if not b_lines:
                continue

            inst = b_lines[0]
            degree = b_lines[1] if len(b_lines) > 1 else ""
            dur = ""
            loc = ""

            for l in b_lines:
                d = re.search(r'(\d{2}/\d{4}\s*[–\-]\s*\d{2}/\d{4}|\d{4}\s*[–\-]\s*(?:Present|\d{4}))', l)
                if d and not dur:
                    dur = d.group(0)
                m_loc = re.search(r'\|\s*([A-Za-z\s,]+)$', l)
                if m_loc and not loc:
                    loc = m_loc.group(1).strip()

            entries.append({
                "institute": inst.rstrip(",| "),
                "university": inst.rstrip(",| "),
                "degree": degree.rstrip(",| "),
                "specialization": degree.rstrip(",| "),
                "cgpa": "",
                "duration": dur,
                "location": loc
            })

        return entries

    @classmethod
    def extract_experience(cls, lines: list) -> list:
        if not lines:
            return []
        entries = []
        curr = None
        bullet_starts = ("worked", "built", "developed", "contributed", "gained", "implemented", "designed", "created", "responsible", "managed", "led", "•", "-", "*", ">", "✓", "○")

        for line in lines:
            l_lower = line.lower().strip()
            is_bullet = l_lower.startswith(bullet_starts) or line.startswith(("•", "-", "*", ">", "✓", "○"))
            is_header = not is_bullet and (re.search(r'\b(?:intern|developer|engineer|manager|lead|architect|analyst|associate)\b', line, re.IGNORECASE) or re.search(r'\d{2}/\d{4}|\d{4}\s*[–\-]', line))

            if is_header:
                if curr:
                    entries.append(curr)
                curr = {
                    "company": line,
                    "role": line,
                    "company_role": line,
                    "employment_type": "Full-time",
                    "duration": "",
                    "location": "",
                    "responsibilities": [],
                    "technologies": []
                }
                dur = re.search(r'(\d{2}/\d{4}\s*[–\-]\s*\d{2}/\d{4}|\d{4}\s*[–\-]\s*(?:Present|\d{4}))', line)
                if dur:
                    curr["duration"] = dur.group(0)
            else:
                if not curr:
                    curr = {"company": "Company", "role": "Experience Entry", "company_role": "Experience Entry", "employment_type": "Full-time", "duration": "", "location": "", "responsibilities": [], "technologies": []}
                clean_resp = re.sub(r'^[•\-\*\>\✓\○\s]+', '', line).strip()
                if clean_resp:
                    curr["responsibilities"].append(clean_resp)

        if curr:
            entries.append(curr)

        return entries

    @classmethod
    def extract_internships(cls, lines: list) -> list:
        return cls.extract_experience(lines)

    @classmethod
    def extract_projects(cls, lines: list) -> list:
        """Parses projects into unified objects with name, description, features, responsibilities, technologies."""
        if not lines:
            return []
        projects = []
        curr = None
        bullet_starts = ("•", "-", "*", ">", "✓", "○", "developed", "built", "implemented", "designed", "created", "technologies:", "tech stack:", "built using:", "features:")

        for line in lines:
            l_str = line.strip()
            if not l_str:
                continue
            l_lower = l_str.lower()

            is_bullet = any(l_lower.startswith(b) for b in bullet_starts) or line.startswith(("•", "-", "*", ">", "✓", "○"))
            is_tech_line = any(k in l_lower for k in ["technologies:", "tech stack:", "built using:"])
            is_feature_line = "features:" in l_lower or "key features:" in l_lower

            if is_tech_line and curr:
                tech_str = re.sub(r'^(?:technologies|tech stack|built using)[:\s]*', '', l_str, flags=re.IGNORECASE)
                toks = [t.strip() for t in re.split(r'[,•\|\/\*\:\(\)]', tech_str) if t.strip() and t.strip().lower() not in ["technologies", "tech stack", "built using"]]
                curr["technologies"].extend(toks)
            elif is_feature_line and curr:
                feat_str = re.sub(r'^(?:features|key features)[:\s]*', '', l_str, flags=re.IGNORECASE)
                toks = [t.strip() for t in re.split(r'[,•\|\/\*\:]', feat_str) if t.strip() and t.strip().lower() != "features"]
                curr["features"].extend(toks)
            elif is_bullet and curr:
                clean_text = re.sub(r'^[•\-\*\>\✓\○\s]+', '', l_str).strip()
                if clean_text.lower().startswith("features:"):
                    feat_str = re.sub(r'^features:[\s]*', '', clean_text, flags=re.IGNORECASE)
                    toks = [t.strip() for t in re.split(r'[,•\|\/\*\:]', feat_str) if t.strip() and t.strip().lower() != "features"]
                    curr["features"].extend(toks)
                elif clean_text.lower().startswith("technologies:"):
                    tech_str = re.sub(r'^technologies:[\s]*', '', clean_text, flags=re.IGNORECASE)
                    toks = [t.strip() for t in re.split(r'[,•\|\/\*\:\(\)]', tech_str) if t.strip() and t.strip().lower() != "technologies"]
                    curr["technologies"].extend(toks)
                else:
                    curr["responsibilities"].append(clean_text)
            else:
                # Project Title line
                if curr:
                    projects.append(curr)
                curr = {
                    "name": l_str,
                    "description": l_str,
                    "features": [],
                    "responsibilities": [],
                    "technologies": [],
                    "github": "",
                    "live_url": "",
                    "duration": ""
                }

        if curr:
            projects.append(curr)

        # Clean duplicates inside projects
        for p in projects:
            p["features"] = sorted(list(set(p["features"])))
            p["technologies"] = sorted(list(set(p["technologies"])))
            p["responsibilities"] = list(dict.fromkeys(p["responsibilities"]))

        return projects

    @classmethod
    def extract_skills(cls, lines: list, text: str) -> dict:
        categorized = {key: [] for key in cls.SKILL_TAXONOMY}
        all_skills_found = {}
        text_lower = text.lower()

        for category, skill_list in cls.SKILL_TAXONOMY.items():
            for skill in skill_list:
                pattern = r'(?<![a-zA-Z0-9])' + re.escape(skill) + r'(?![a-zA-Z0-9])'
                if re.search(pattern, text_lower):
                    formatted = skill.upper() if len(skill) <= 3 and skill not in ["css", "sql", "git"] else skill.title()
                    if skill == "c++": formatted = "C++"
                    elif skill == "c#": formatted = "C#"
                    elif skill == "node.js": formatted = "Node.js"
                    elif skill == "react.js": formatted = "React.js"
                    elif skill == "next.js": formatted = "Next.js"
                    elif skill == "vue.js": formatted = "Vue.js"
                    elif skill == "express.js": formatted = "Express.js"
                    elif skill == "ci/cd": formatted = "CI/CD"
                    elif skill == "rest api": formatted = "REST API"
                    elif skill == "aws": formatted = "AWS"
                    elif skill == "gcp": formatted = "GCP"

                    categorized[category].append(formatted)
                    all_skills_found[skill.lower()] = formatted

        skills_text = " ".join(lines) if lines else text
        tokens = re.split(r'[,•\|\/\*\:\n]', skills_text)
        skip_words = {"technical", "skills", "experience", "education", "projects", "tools", "proficiencies", "and", "or", "etc", "using", "with", "level", "proficient", "basic", "advanced", "intermediate", "development", "intern", "backend", "frontend", "integration", "deployment", "basics", "filtering", "management", "reports", "authentication", "features"}

        for tok in tokens:
            clean_tok = re.sub(r'[^a-zA-Z0-9#+.\s-]', '', tok).strip()
            tok_lower = clean_tok.lower()
            if 2 <= len(clean_tok) <= 25 and tok_lower not in skip_words and len(clean_tok.split()) <= 2 and not re.search(r'\d{2,}', clean_tok):
                if tok_lower not in all_skills_found:
                    formatted = clean_tok.title() if not clean_tok.isupper() else clean_tok
                    if tok_lower in ["html", "css", "sql", "aws", "gcp", "api", "xml", "json", "rest"]:
                        formatted = clean_tok.upper()
                    categorized["developer_tools"].append(formatted)
                    all_skills_found[tok_lower] = formatted

        for cat in categorized:
            categorized[cat] = sorted(list(set(categorized[cat])))

        return categorized

    @classmethod
    def extract_certifications(cls, lines: list) -> list:
        return lines

    @classmethod
    def extract_languages(cls, lines: list, text: str) -> list:
        found_langs = set()
        common_langs = ["English", "Spanish", "French", "German", "Chinese", "Hindi", "Gujarati", "Japanese", "Arabic", "Portuguese", "Russian", "Italian", "Dutch", "Korean"]

        if lines:
            lang_str = " ".join(lines)
            tokens = re.split(r'[,•\|\/\s]', lang_str)
            for tok in tokens:
                clean_tok = tok.strip().title()
                if clean_tok in common_langs or (len(clean_tok) >= 3 and clean_tok.isalpha() and clean_tok not in ["Fluent", "Native", "Basic", "Proficient", "Intermediate", "Language"]):
                    found_langs.add(clean_tok)

        if not found_langs:
            for lang in common_langs:
                if re.search(r'\b' + lang + r'\b', text, re.IGNORECASE):
                    found_langs.add(lang)

        return sorted(list(found_langs))

    @classmethod
    def extract_courses(cls, lines: list) -> list:
        return lines

    @classmethod
    def extract_achievements(cls, lines: list) -> list:
        return lines

    @classmethod
    def build_json(cls, text: str) -> dict:
        text = cls.clean_text(text)
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        detected_sections = cls.detect_sections(text)

        personal_info = cls.extract_personal_information(lines, text)

        san_edu = cls._sanitize_lines(detected_sections.get("education", []), personal_info)
        san_exp = cls._sanitize_lines(detected_sections.get("experience", []), personal_info)
        san_intern = cls._sanitize_lines(detected_sections.get("internship", []), personal_info)
        san_proj = cls._sanitize_lines(detected_sections.get("projects", []), personal_info)

        skills_categorized = cls.extract_skills(detected_sections.get("skills", []), text)

        flat_skills = []
        for cat_skills in skills_categorized.values():
            flat_skills.extend(cat_skills)
        flat_skills = sorted(list(set(flat_skills)))

        parsed = {
            "personal_information": personal_info,
            "name": personal_info["full_name"],
            "professional_title": personal_info["professional_title"],
            "email": personal_info["email"],
            "phone": personal_info["phone"],
            "location": personal_info["location"],
            "address": personal_info["address"],
            "linkedin": personal_info["linkedin"],
            "github": personal_info["github"],
            "portfolio": personal_info["portfolio"],
            "website": personal_info["website"],
            "dob": personal_info["dob"],
            "nationality": personal_info["nationality"],
            "contact_info": personal_info,
            "summary": " ".join(detected_sections.get("summary", [])) or " ".join(detected_sections.get("objective", [])),
            "objective": " ".join(detected_sections.get("objective", [])),
            "education": cls.extract_education(san_edu),
            "structured_education": cls.extract_education(san_edu),
            "experience": cls.extract_experience(san_exp),
            "structured_experience": cls.extract_experience(san_exp),
            "internships": cls.extract_internships(san_intern),
            "structured_internship": cls.extract_internships(san_intern),
            "projects": cls.extract_projects(san_proj),
            "structured_projects": cls.extract_projects(san_proj),
            "skills": skills_categorized,
            "flat_skills": flat_skills,
            "certifications": cls.extract_certifications(detected_sections.get("certifications", [])),
            "courses": cls.extract_courses(detected_sections.get("courses", [])),
            "achievements": cls.extract_achievements(detected_sections.get("achievements", [])),
            "languages": cls.extract_languages(detected_sections.get("languages", []), text),
            "volunteer": detected_sections.get("volunteer", []),
            "research": detected_sections.get("research", []),
            "publications": detected_sections.get("publications", []),
            "interests": detected_sections.get("interests", []),
            "references": detected_sections.get("references", []),
            "word_count": len(text.split()),
            "character_count": len(text)
        }

        return parsed

    @classmethod
    def parse_resume(cls, text: str) -> dict:
        return cls.build_json(text)

    # Helper methods
    @classmethod
    def _fix_spaced_text(cls, text: str) -> str:
        if not text:
            return ""

        def unspace_match(match):
            return match.group(0).replace(" ", "")

        fixed = re.sub(r'(?<![A-Za-z0-9])(?:[A-Za-z]\s+){1,}[A-Za-z](?![A-Za-z0-9])', unspace_match, text)
        return fixed

    @classmethod
    def _extract_pdf(cls, path: Path) -> str:
        text = ""
        try:
            import pypdf
            reader = pypdf.PdfReader(str(path))
            for page in reader.pages:
                extracted = page.extract_text()
                if extracted:
                    text += extracted + "\n"
        except Exception as e:
            logger.warning(f"pypdf extraction error on {path}: {e}")

        if len(text.strip()) < 30:
            try:
                import PyPDF2
                with open(path, "rb") as f:
                    reader = PyPDF2.PdfReader(f)
                    pypdf_text = ""
                    for page in reader.pages:
                        t = page.extract_text()
                        if t:
                            pypdf_text += t + "\n"
                    if len(pypdf_text.strip()) > len(text.strip()):
                        text = pypdf_text
            except Exception as e:
                logger.warning(f"PyPDF2 extraction error on {path}: {e}")

        return text

    @classmethod
    def _extract_docx(cls, path: Path) -> str:
        full_text = []

        try:
            import docx
            doc = docx.Document(path)
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text.strip())
            for table in doc.tables:
                for row in table.rows:
                    row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_cells:
                        full_text.append(" • ".join(row_cells))
        except Exception as e:
            logger.warning(f"docx python-docx extraction warning on {path}: {e}")

        return "\n".join(full_text)

    @classmethod
    def _extract_image_ocr(cls, path: Path) -> str:
        text = ""
        try:
            from PIL import Image
            import pytesseract
            img = Image.open(path)
            text = pytesseract.image_to_string(img)
        except Exception as e:
            logger.warning(f"pytesseract OCR warning on {path}: {e}")

        if not text.strip():
            text = f"Image Resume ({path.name}) loaded."
        return text

    @classmethod
    def detect_sections(cls, text: str) -> dict:
        lines = [l.strip() for l in text.splitlines() if l.strip()]
        sections = {key: [] for key in cls.SECTION_MAP}
        sections["header_unassigned"] = []

        current_section = None

        for line in lines:
            normalized = cls.normalize_heading(line)

            if current_section == "skills" and not normalized:
                if re.search(r'\b(?:intern|developer|engineer|manager|worked|built|developed|contributed)\b', line, re.IGNORECASE) or re.search(r'\d{2}/\d{4}', line):
                    normalized = "experience"

            if normalized:
                current_section = normalized
            elif current_section:
                sections[current_section].append(line)
            else:
                sections["header_unassigned"].append(line)

        return sections

    @staticmethod
    def _extract_email(text: str) -> str:
        match = re.search(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}', text)
        return match.group(0) if match else ""

    @staticmethod
    def _extract_phone(text: str) -> str:
        matches = re.findall(r'(?:\+?\d{1,3}[-.\s]?)?\(?\d{3,5}\)?[-.\s]?\d{3,5}[-.\s]?\d{3,5}', text)
        for m in matches:
            digits = re.sub(r'\D', '', m)
            if 7 <= len(digits) <= 15 and not m.startswith("202") and not m.startswith("199"):
                return m.strip()
        return ""

    @staticmethod
    def _extract_links(text: str) -> dict:
        links = {"linkedin": "", "github": "", "portfolio": ""}

        linkedin = re.search(r'(https?://)?(www\.)?linkedin\.com/in/[\w-]+', text, re.IGNORECASE)
        if linkedin:
            links["linkedin"] = linkedin.group(0)

        github = re.search(r'(https?://)?(www\.)?github\.com/[\w-]+', text, re.IGNORECASE)
        if github:
            links["github"] = github.group(0)

        portfolio = re.search(r'(https?://)?(www\.)?[\w-]+\.(io|dev|me|tech|site|com|net|org)', text, re.IGNORECASE)
        if portfolio and not linkedin and not github:
            links["portfolio"] = portfolio.group(0)

        return links

    @staticmethod
    def _extract_name(lines: list, email: str) -> str:
        if not lines:
            return "Candidate"

        skip_words = ["resume", "curriculum", "vitae", "cv", "profile", "contact", "summary", "address", "phone", "email", "institute", "university", "college", "school", "academy", "education", "bachelor", "master", "secondary", "certificate", "senior"]

        for line in lines[:6]:
            clean_line = line.strip()
            clean_lower = clean_line.lower()

            if any(kw in clean_lower for kw in skip_words):
                continue
            if "@" in clean_line or "http" in clean_line or "www." in clean_line:
                continue

            letters_only = re.sub(r'[^a-zA-Z\s]', '', clean_line).strip()
            words = letters_only.split()

            if 1 <= len(words) <= 4 and len(clean_line) < 40:
                return clean_line.title()

        if email:
            prefix = email.split("@")[0]
            clean_prefix = re.sub(r'[^a-zA-Z.]', ' ', prefix).replace(".", " ").strip()
            if clean_prefix:
                return clean_prefix.title()

        return "Candidate"

    @staticmethod
    def _extract_professional_title(lines: list, text: str) -> str:
        match = re.search(r'\b(Python\s+Developer|Full\s+Stack\s+Developer|Frontend\s+Developer|Backend\s+Developer|Software\s+Engineer|Data\s+Scientist|UI/UX\s+Designer)\b', text, re.IGNORECASE)
        if match:
            return match.group(0).title()
        return "Software Developer"

    @staticmethod
    def _extract_location(text: str) -> str:
        match = re.search(r'([A-Z][a-zA-Z\s]{2,18},\s*[A-Z]{2,3}(?:,\s*[A-Z][a-zA-Z\s]{2,15})?|[A-Z][a-zA-Z\s]{2,18},\s*[A-Z][a-zA-Z\s]{2,15})', text)
        if match and not any(kw in match.group(0).lower() for kw in ['developer', 'engineer', 'intern', 'self']):
            return match.group(0)
        return ""
